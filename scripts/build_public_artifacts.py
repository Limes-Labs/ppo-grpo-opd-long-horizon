#!/usr/bin/env python3
"""Build public PDF/DOCX paper artifacts from canonical experiment JSON."""

from __future__ import annotations

import argparse
import hashlib
import json
import subprocess
import sys
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any


TITLE = "PPO, GRPO, and On-Policy Distillation for Long-Horizon Post-Training"
SUBTITLE = "A Limes Labs public research report on temporal credit assignment"
RUNTIME_NOTE = "Generated from canonical JSON experiment artifacts."
GLM_SOURCE_URL = "https://z.ai/blog/glm-5.2"
REQUIRED_TEXT_PHRASES = [
    "PPO, GRPO",
    "16 clear",
    "near tie",
    "Z.ai",
    "not independent causal evidence",
]


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text())


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def git_value(args: list[str], cwd: Path) -> str:
    try:
        return subprocess.check_output(["git", *args], cwd=cwd, text=True).strip()
    except Exception:
        return "unknown"


def fmt(value: float, digits: int = 3) -> str:
    return f"{value:.{digits}f}"


def text_check_payload(text: str, required_phrases: list[str]) -> dict[str, Any]:
    return {
        "required_phrases": required_phrases,
        "missing_phrases": [phrase for phrase in required_phrases if phrase not in text],
        "character_count": len(text),
    }


def top_cases(result: dict[str, Any], limit: int = 8) -> list[dict[str, Any]]:
    return sorted(
        result["cases"],
        key=lambda case: abs(case["mean_critic_minus_group_correlation"]),
        reverse=True,
    )[:limit]


def report_axis(axis: str) -> str:
    return {
        "counterexample": "counter",
        "critic_budget": "budget",
        "group_size": "group",
        "observability": "observe",
        "sparse_reward": "sparse",
    }.get(axis, axis)


def create_chart_pngs(result: dict[str, Any], output_dir: Path) -> list[Path]:
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError as exc:
        raise RuntimeError("Pillow is required to generate PNG charts") from exc

    output_dir.mkdir(parents=True, exist_ok=True)

    def font(size: int, bold: bool = False) -> Any:
        candidates = [
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
            "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
        ]
        for candidate in candidates:
            try:
                return ImageFont.truetype(candidate, size)
            except OSError:
                continue
        return ImageFont.load_default()

    title_font = font(34, bold=True)
    body_font = font(20)
    small_font = font(16)
    cases = result["cases"]

    delta_path = output_dir / "deep_matrix_delta.png"
    width = 1700
    row_h = 52
    top = 128
    left = 470
    right = 100
    height = top + row_h * len(cases) + 88
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((44, 32), "Critic minus group correlation by case", font=title_font, fill="#0B2545")
    draw.text(
        (44, 78),
        "Positive values favor critic-style TD; negative values favor group-relative advantage.",
        font=body_font,
        fill="#333333",
    )
    chart_w = width - left - right
    zero_x = left + chart_w // 2
    max_abs = max(0.1, max(abs(case["mean_critic_minus_group_correlation"]) for case in cases))
    draw.line((zero_x, top - 24, zero_x, height - 60), fill="#666666", width=2)
    for idx, case in enumerate(cases):
        y = top + idx * row_h
        delta = case["mean_critic_minus_group_correlation"]
        bar_len = int(abs(delta) / max_abs * (chart_w / 2 - 24))
        color = "#1F6F4A" if delta >= 0 else "#9B1C1C"
        x0 = zero_x if delta >= 0 else zero_x - bar_len
        draw.rounded_rectangle((x0, y + 8, x0 + bar_len, y + 34), radius=6, fill=color)
        draw.text((44, y + 8), case["case_name"], font=small_font, fill="#111111")
        text_x = zero_x + bar_len + 10 if delta >= 0 else zero_x - bar_len - 82
        draw.text((text_x, y + 10), f"{delta:+.3f}", font=small_font, fill=color)
    image.save(delta_path)

    coverage_path = output_dir / "deep_matrix_coverage.png"
    width = 1300
    height = 900
    margin = 112
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 34), "Critic coverage vs estimator advantage", font=title_font, fill="#0B2545")
    draw.text(
        (60, 82),
        "A group-relative counterexample appears when critic state coverage is weak.",
        font=body_font,
        fill="#333333",
    )
    plot_left = margin
    plot_top = 150
    plot_w = width - 2 * margin
    plot_h = height - 260
    draw.rectangle((plot_left, plot_top, plot_left + plot_w, plot_top + plot_h), outline="#AAB4C0", width=2)
    deltas = [case["mean_critic_minus_group_correlation"] for case in cases]
    min_delta = min(-0.08, min(deltas))
    max_delta = max(0.08, max(deltas))

    def x_of(hit: float) -> int:
        return int(plot_left + hit * plot_w)

    def y_of(delta: float) -> int:
        return int(plot_top + (max_delta - delta) / (max_delta - min_delta) * plot_h)

    zero_y = y_of(0.0)
    draw.line((plot_left, zero_y, plot_left + plot_w, zero_y), fill="#777777", width=2)
    draw.text((plot_left + plot_w // 2 - 122, height - 78), "critic exact-state hit rate", font=body_font, fill="#333333")
    draw.text((28, plot_top + plot_h // 2), "delta r", font=body_font, fill="#333333")
    for tick in [0.0, 0.25, 0.5, 0.75, 1.0]:
        x = x_of(tick)
        draw.line((x, plot_top + plot_h, x, plot_top + plot_h + 10), fill="#777777", width=1)
        draw.text((x - 16, plot_top + plot_h + 18), f"{tick:.2f}", font=small_font, fill="#444444")
    for case in cases:
        x = x_of(case["mean_critic_exact_state_rate"])
        y = y_of(case["mean_critic_minus_group_correlation"])
        color = "#1F6F4A" if case["mean_critic_minus_group_correlation"] >= 0 else "#9B1C1C"
        draw.ellipse((x - 9, y - 9, x + 9, y + 9), fill=color, outline="#111111")
        if case["winner_by_mean_correlation"] == "group":
            draw.text((x + 14, y - 16), "group wins", font=small_font, fill=color)
    image.save(coverage_path)

    return [delta_path, coverage_path]


def build_pdf(result: dict[str, Any], charts: list[Path], output: Path) -> None:
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            Image,
            PageBreak,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except ImportError as exc:
        raise RuntimeError("reportlab is required to generate the PDF") from exc

    output.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="TitleCenter",
            parent=styles["Title"],
            alignment=TA_CENTER,
            fontSize=22,
            leading=27,
            textColor=colors.HexColor("#0B2545"),
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SectionHeading",
            parent=styles["Heading1"],
            textColor=colors.HexColor("#2E74B5"),
            fontSize=15,
            leading=19,
            spaceBefore=14,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyTight",
            parent=styles["BodyText"],
            fontSize=9.5,
            leading=12,
            spaceAfter=6,
        )
    )

    doc = SimpleDocTemplate(
        str(output),
        pagesize=letter,
        rightMargin=0.78 * inch,
        leftMargin=0.78 * inch,
        topMargin=0.72 * inch,
        bottomMargin=0.72 * inch,
    )
    story: list[Any] = []
    story.append(Paragraph(TITLE, styles["TitleCenter"]))
    story.append(Paragraph(SUBTITLE, styles["Heading2"]))
    story.append(Paragraph("Limes Labs Research Workstream | Public draft | 2026-06-18", styles["BodyTight"]))
    story.append(Spacer(1, 10))
    story.append(
        Paragraph(
            "Abstract. We compare PPO-style critic/value-model advantage estimation, GRPO-style group-relative advantages, and OPD/OPSD distillation for long-horizon post-training. In a multi-seed toy matrix with known oracle advantages, critic-style TD has the higher mean correlation in 17/18 fixed regimes; a confidence-interval reading gives 16 clear critic-favorable cases, 1 near tie, and 1 clear group-favorable counterexample. The conclusion is conditional: PPO-style critics look better when temporal state information is observable and learnable; GRPO remains attractive when reward contrast is reliable and critic coverage/cost is poor.",
            styles["BodyTight"],
        )
    )

    story.append(Paragraph("Key Results", styles["SectionHeading"]))
    summary = result["overall"]
    key_table = Table(
        [
            ["Seeds", "Cases", "Clear critic", "Near tie", "Clear group", "Mean delta r"],
            [
                str(result["seed_count"]),
                str(result["case_count"]),
                str(summary["clear_critic_cases_by_ci95"]),
                str(summary["near_tie_cases_by_ci95"]),
                str(summary["clear_group_cases_by_ci95"]),
                fmt(summary["mean_delta_correlation"]),
            ],
        ],
        colWidths=[0.75 * inch, 0.65 * inch, 1.0 * inch, 0.85 * inch, 0.95 * inch, 1.1 * inch],
    )
    key_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#9AA6B2")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.append(key_table)
    story.append(Spacer(1, 8))

    for chart in charts:
        story.append(Image(str(chart), width=6.5 * inch, height=3.65 * inch))
        story.append(Spacer(1, 10))

    story.append(PageBreak())
    story.append(Paragraph("Methods", styles["SectionHeading"]))
    story.append(
        Paragraph(
            "The toy environment has variable-length trajectories with help, harm, and wait actions. A terminal verifier supplies a sparse success reward; known finite-state dynamics provide an oracle state-action advantage. The group-relative estimator normalizes terminal returns within a prompt group and broadcasts the scalar to each token. The critic estimator fits a tabular value model from returns-to-go and computes one-step TD advantages. This measures estimator fidelity, not closed-loop PPO or GRPO learning.",
            styles["BodyTight"],
        )
    )
    story.append(Paragraph("Deep Matrix Results", styles["SectionHeading"]))
    rows = [["Case", "Axis", "Mean", "CI read", "Group r", "Critic r", "Delta r", "95% CI"]]
    for case in result["cases"]:
        rows.append(
            [
                case["case_name"],
                report_axis(case["axis"]),
                case["winner_by_mean_correlation"],
                case["evidence_by_ci95"],
                fmt(case["mean_group_correlation"]),
                fmt(case["mean_critic_correlation"]),
                fmt(case["mean_critic_minus_group_correlation"]),
                "+/- " + fmt(case["ci95_critic_minus_group_correlation"]),
            ]
        )
    table = Table(
        rows,
        repeatRows=1,
        colWidths=[
            1.7 * inch,
            0.72 * inch,
            0.47 * inch,
            0.7 * inch,
            0.56 * inch,
            0.56 * inch,
            0.56 * inch,
            0.6 * inch,
        ],
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B8C0CC")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 6.8),
                ("ALIGN", (2, 1), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(table)

    story.append(PageBreak())
    story.append(Paragraph("Interpretation", styles["SectionHeading"]))
    for paragraph in [
        "The evidence supports the PPO-favorable hypothesis in this controlled toy: when progress state is observable and value coverage is adequate, critic-style TD advantages track oracle token credit far better than response-level group scalars.",
        "The counterexample matters. When the critic is blind and undercovered, terminal group-relative outcome information can be more useful than a poor value table. The scientific claim is therefore not that PPO is universally better; it is that value-based temporal information becomes valuable in the long-horizon regime if it can be learned.",
        "Z.ai's GLM-5.2 report is a useful industry case study because it describes compacted, variable-length long-horizon agent rollouts and reports moving from group-wise optimization to critic-based PPO with token-level advantages. It is not independent causal evidence for PPO over GRPO; it is a motivating example that matches the failure mode studied here.",
    ]:
        story.append(Paragraph(paragraph, styles["BodyTight"]))

    story.append(Paragraph("Cost, Safety, and Caveats", styles["SectionHeading"]))
    for paragraph in [
        "Fair comparisons must charge critic memory and training, reward/verifier calls, group sampling, teacher/self-teacher compute, anti-hacking filters, LLM judges, sandboxing, blocked-call handling, and trajectory storage.",
        "Coding-agent RL with pass/fail rewards can incentivize shortcuts. Online tool-call monitoring and invalid-action handling should be treated as part of the training protocol, not merely post-hoc evaluation.",
        "The toy has an oracle and tabular critic. Real systems add neural approximation, KL controllers, stale policies, reward-model error, tool failures, and distributed rollout infrastructure. Stronger GRPO variants and process rewards remain future work.",
    ]:
        story.append(Paragraph(paragraph, styles["BodyTight"]))

    story.append(Paragraph("Selected References", styles["SectionHeading"]))
    refs = [
        "[1] Schulman et al. Proximal Policy Optimization Algorithms. 2017.",
        "[5] Shao et al. DeepSeekMath: Group Relative Policy Optimization. 2024.",
        "[6] DeepSeek-AI. DeepSeek-R1. 2025.",
        "[13] Zhao et al. Self-Distilled Reasoner: OPSD. 2026.",
        "[19] Z.ai. GLM-5.2: Built for Long-Horizon Tasks. 2026. https://z.ai/blog/glm-5.2",
    ]
    for ref in refs:
        story.append(Paragraph(ref, styles["BodyTight"]))

    def page(canvas: Any, doc_obj: Any) -> None:
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#666666"))
        canvas.drawString(0.78 * inch, 0.42 * inch, "Limes Labs | PPO-GRPO-OPD Long-Horizon Workstream")
        canvas.drawRightString(7.72 * inch, 0.42 * inch, f"Page {doc_obj.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=page, onLaterPages=page)


def build_docx(result: dict[str, Any], charts: list[Path], output: Path) -> None:
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("python-docx is required to generate the DOCX") from exc

    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(SUBTITLE).italic = True
    doc.add_paragraph("Limes Labs Research Workstream | Public draft | 2026-06-18")

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "We compare PPO-style critic/value-model advantage estimation, GRPO-style group-relative advantages, and OPD/OPSD distillation for long-horizon post-training. In a multi-seed toy matrix with known oracle advantages, critic-style TD has the higher mean correlation in 17/18 fixed regimes; a confidence-interval reading gives 16 clear critic-favorable cases, 1 near tie, and 1 clear group-favorable counterexample. The conclusion is conditional: PPO-style critics look better when temporal state information is observable and learnable; GRPO remains attractive when reward contrast is reliable and critic coverage/cost is poor."
    )

    doc.add_heading("Key Results", level=1)
    key_table = doc.add_table(rows=2, cols=6)
    key_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Seeds", "Cases", "Clear critic", "Near tie", "Clear group", "Mean delta r"]
    values = [
        str(result["seed_count"]),
        str(result["case_count"]),
        str(result["overall"]["clear_critic_cases_by_ci95"]),
        str(result["overall"]["near_tie_cases_by_ci95"]),
        str(result["overall"]["clear_group_cases_by_ci95"]),
        fmt(result["overall"]["mean_delta_correlation"]),
    ]
    for idx, text in enumerate(headers):
        key_table.cell(0, idx).text = text
        key_table.cell(1, idx).text = values[idx]
    for row in key_table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for chart in charts:
        doc.add_picture(str(chart), width=Inches(6.4))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Methods", level=1)
    doc.add_paragraph(
        "The toy environment has variable-length trajectories with help, harm, and wait actions. A terminal verifier supplies a sparse success reward; known finite-state dynamics provide an oracle state-action advantage. The group-relative estimator normalizes terminal returns within a prompt group and broadcasts the scalar to every token. The critic estimator fits a tabular value model from returns-to-go and computes one-step TD advantages."
    )

    doc.add_heading("Deep Matrix Results", level=1)
    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Case", "Axis", "Mean", "CI read", "Group r", "Critic r", "Delta r", "95% CI"]
    for idx, text in enumerate(headers):
        table.cell(0, idx).text = text
    for case in result["cases"]:
        row = table.add_row().cells
        values = [
            case["case_name"],
            report_axis(case["axis"]),
            case["winner_by_mean_correlation"],
            case["evidence_by_ci95"],
            fmt(case["mean_group_correlation"]),
            fmt(case["mean_critic_correlation"]),
            fmt(case["mean_critic_minus_group_correlation"]),
            "+/- " + fmt(case["ci95_critic_minus_group_correlation"]),
        ]
        for idx, text in enumerate(values):
            row[idx].text = text

    doc.add_heading("Interpretation", level=1)
    for text in [
        "The evidence supports the PPO-favorable hypothesis in this controlled toy: when progress state is observable and value coverage is adequate, critic-style TD advantages track oracle token credit far better than response-level group scalars.",
        "The counterexample matters. When the critic is blind and undercovered, terminal group-relative outcome information can be more useful than a poor value table.",
        "Z.ai's GLM-5.2 report is a useful industry case study because it describes compacted, variable-length long-horizon agent rollouts and reports moving from group-wise optimization to critic-based PPO with token-level advantages. It is not independent causal evidence for PPO over GRPO.",
    ]:
        doc.add_paragraph(text)

    doc.add_heading("Cost, Safety, and Caveats", level=1)
    for text in [
        "Fair comparisons must charge critic memory and training, reward/verifier calls, group sampling, teacher/self-teacher compute, anti-hacking filters, LLM judges, sandboxing, blocked-call handling, and trajectory storage.",
        "Coding-agent RL with pass/fail rewards can incentivize shortcuts. Online tool-call monitoring and invalid-action handling should be treated as part of the training protocol, not merely post-hoc evaluation.",
        "The toy has an oracle and tabular critic. Real systems add neural approximation, KL controllers, stale policies, reward-model error, tool failures, and distributed rollout infrastructure.",
    ]:
        doc.add_paragraph(text)

    doc.add_heading("Selected References", level=1)
    for ref in [
        "[1] Schulman et al. Proximal Policy Optimization Algorithms. 2017.",
        "[5] Shao et al. DeepSeekMath: Group Relative Policy Optimization. 2024.",
        "[6] DeepSeek-AI. DeepSeek-R1. 2025.",
        "[13] Zhao et al. Self-Distilled Reasoner: OPSD. 2026.",
        "[19] Z.ai. GLM-5.2: Built for Long-Horizon Tasks. 2026. " + GLM_SOURCE_URL,
    ]:
        doc.add_paragraph(ref)

    footer = section.footer.paragraphs[0]
    footer.text = "Limes Labs | PPO-GRPO-OPD Long-Horizon Workstream"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(output)


def verify_pdf_artifact(path: Path) -> dict[str, Any]:
    payload: dict[str, Any] = {
        "generated": path.exists() and path.stat().st_size > 0,
        "bytes": path.stat().st_size if path.exists() else 0,
    }
    try:
        from pypdf import PdfReader
    except ImportError:
        payload["text_check"] = "not_run_pypdf_unavailable"
        return payload

    reader = PdfReader(str(path))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    payload["page_count"] = len(reader.pages)
    payload["text_check"] = text_check_payload(text, REQUIRED_TEXT_PHRASES)
    return payload


def verify_docx_artifact(path: Path) -> dict[str, Any]:
    payload: dict[str, Any] = {
        "generated": path.exists() and path.stat().st_size > 0,
        "bytes": path.stat().st_size if path.exists() else 0,
    }
    try:
        from docx import Document
    except ImportError:
        payload["text_check"] = "not_run_python_docx_unavailable"
        return payload

    document = Document(str(path))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    with zipfile.ZipFile(path) as archive:
        media = [
            name
            for name in archive.namelist()
            if name.startswith("word/media/") and not name.endswith("/")
        ]
    payload["paragraph_count"] = len(document.paragraphs)
    payload["table_count"] = len(document.tables)
    payload["inline_shape_count"] = len(document.inline_shapes)
    payload["media_count"] = len(media)
    payload["text_check"] = text_check_payload(text, REQUIRED_TEXT_PHRASES)
    payload["visual_render"] = "not_run_soffice_unavailable"
    return payload


def build_manifest(
    *,
    repo: Path,
    inputs: list[Path],
    outputs: list[Path],
    render_checks: dict[str, Any],
) -> dict[str, Any]:
    return {
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "generator": "scripts/build_public_artifacts.py",
        "artifact_scope": "abridged public PDF/DOCX report derived from the canonical matrix, not a full PAPER.md export",
        "git_commit": git_value(["rev-parse", "HEAD"], repo),
        "git_status_short": git_value(["status", "--short"], repo),
        "git_context_note": (
            "The manifest records the worktree at generation time. When generated "
            "artifacts are committed together with this manifest, the final commit "
            "will necessarily differ; rebuild from a clean checkout to refresh it."
        ),
        "source_urls": {
            "z_ai_glm_5_2": GLM_SOURCE_URL,
        },
        "inputs": {str(path): sha256_file(path) for path in inputs},
        "outputs": {
            str(path): {
                "sha256": sha256_file(path),
                "bytes": path.stat().st_size,
            }
            for path in outputs
        },
        "render_checks": render_checks,
    }


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--matrix-json", type=Path, default=Path("results/deep_matrix_20seed.json"))
    parser.add_argument("--public-dir", type=Path, default=Path("public"))
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    repo = Path.cwd()
    result = read_json(args.matrix_json)
    public_dir = args.public_dir
    figures_dir = public_dir / "figures"
    pdf = public_dir / "ppo_grpo_opd_long_horizon.pdf"
    docx = public_dir / "ppo_grpo_opd_long_horizon.docx"
    manifest_path = public_dir / "artifact_manifest.json"

    charts = create_chart_pngs(result, figures_dir)
    build_pdf(result, charts, pdf)
    build_docx(result, charts, docx)

    render_checks = {
        "pdf": verify_pdf_artifact(pdf),
        "docx": verify_docx_artifact(docx),
    }
    outputs = [pdf, docx, *charts]
    manifest = build_manifest(
        repo=repo,
        inputs=[args.matrix_json],
        outputs=outputs,
        render_checks=render_checks,
    )
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(json.dumps(manifest, allow_nan=False, indent=2, sort_keys=True) + "\n")

    print(f"wrote {pdf}")
    print(f"wrote {docx}")
    for chart in charts:
        print(f"wrote {chart}")
    print(f"wrote {manifest_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
